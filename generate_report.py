"""
Generate the complete university-level project report as a .docx file.
Run: python3 generate_report.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── helpers ────────────────────────────────────────────────────────────────

def set_font(run, name="Times New Roman", size=12, bold=False,
             italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def para_fmt(para, space_before=0, space_after=8,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=None):
    fmt = para.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after  = Pt(space_after)
    fmt.alignment    = align
    if line_spacing:
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = Pt(line_spacing)

def add_heading(doc, text, level=1, numbered="", space_before=18):
    para = doc.add_paragraph()
    para_fmt(para, space_before=space_before, space_after=6,
             align=WD_ALIGN_PARAGRAPH.LEFT)
    prefix = f"{numbered}  " if numbered else ""
    run = para.add_run(prefix + text)
    if level == 0:   # Chapter
        set_font(run, size=15, bold=True, color=(0, 0, 0))
    elif level == 1: # Section
        set_font(run, size=13, bold=True)
    elif level == 2: # Sub-section
        set_font(run, size=12, bold=True)
    else:
        set_font(run, size=12, bold=False, italic=True)
    return para

def add_body(doc, text, bold_spans=None):
    para = doc.add_paragraph()
    para_fmt(para, space_before=2, space_after=8)
    run = para.add_run(text)
    set_font(run)
    return para

def add_bullet(doc, items, level=0):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para_fmt(para, space_before=1, space_after=3,
                 align=WD_ALIGN_PARAGRAPH.LEFT)
        run = para.add_run(item)
        set_font(run, size=11)

def add_numbered_list(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Number")
        para_fmt(para, space_before=1, space_after=3,
                 align=WD_ALIGN_PARAGRAPH.LEFT)
        run = para.add_run(item)
        set_font(run, size=11)

def add_table(doc, headers, rows, caption=""):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=10, bold=True, color=(255, 255, 255))
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F3864")
        cell._tc.get_or_add_tcPr().append(shd)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = "DEEAF1" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_font(run, size=10)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            cell._tc.get_or_add_tcPr().append(shd)
    # Caption
    if caption:
        cp = doc.add_paragraph()
        para_fmt(cp, space_before=4, space_after=12,
                 align=WD_ALIGN_PARAGRAPH.CENTER)
        run = cp.add_run(caption)
        set_font(run, size=10, italic=True)
    return table

def add_code_block(doc, code_text):
    for line in code_text.strip().splitlines():
        para = doc.add_paragraph()
        para_fmt(para, space_before=0, space_after=0,
                 align=WD_ALIGN_PARAGRAPH.LEFT)
        run = para.add_run(line if line.strip() else " ")
        set_font(run, name="Courier New", size=9)
        para.paragraph_format.left_indent = Cm(1)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        para._p.get_or_add_pPr().append(shd)

def add_page_break(doc):
    doc.add_page_break()

def add_caption(doc, text):
    p = doc.add_paragraph()
    para_fmt(p, space_before=2, space_after=10,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(text)
    set_font(r, size=10, italic=True, color=(89, 89, 89))

def rule(doc):
    p = doc.add_paragraph()
    para_fmt(p, space_before=2, space_after=2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFBFBF")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ═══════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ═══════════════════════════════════════════════════════════════
doc = Document()

# ── page margins ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── default paragraph style ───────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════
def add_centered(doc, text, size=12, bold=False, italic=False,
                 color=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    para_fmt(p, space_before=space_before, space_after=space_after,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic, color=color)
    return p

add_centered(doc, "FATIMA JINNAH WOMEN UNIVERSITY",
             size=16, bold=True, color=(31, 56, 100),
             space_before=10, space_after=4)
add_centered(doc, "Department of Software Engineering",
             size=13, bold=True, color=(31, 56, 100), space_after=4)
rule(doc)
add_centered(doc, " ", size=10, space_after=2)
add_centered(doc, "🔐", size=36, space_after=4)
add_centered(doc, "Information Security", size=12, italic=True,
             color=(89, 89, 89), space_after=4)
add_centered(doc, "PROJECT REPORT", size=14, bold=True, space_after=4)
rule(doc)
add_centered(doc, " ", size=10, space_after=2)
add_centered(doc, "Cryptanalysis of Weak Password Hashing Systems",
             size=16, bold=True, color=(31, 56, 100), space_after=4)
add_centered(doc, "and AI-Based Defence Mechanism",
             size=16, bold=True, color=(31, 56, 100), space_after=20)
add_centered(doc, " ", size=10, space_after=4)

# Submission table
tbl = doc.add_table(rows=4, cols=2)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
meta = [
    ("Submitted To:", "Mukhtiar Bano"),
    ("Submitted By:", "Amina Noor (2023-BSE-007)\nHamail Fatima (2023-BSE-023)\nHajra Sarwar (2023-BSE-022)"),
    ("Course:", "Information Security"),
    ("Date:", "22 May 2026"),
]
for i, (label, val) in enumerate(meta):
    r0 = tbl.rows[i].cells[0].paragraphs[0]
    r1 = tbl.rows[i].cells[1].paragraphs[0]
    r0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rl = r0.add_run(label)
    set_font(rl, size=11, bold=True)
    rv = r1.add_run(val)
    set_font(rv, size=11)
    fill = "DEEAF1" if i % 2 == 0 else "FFFFFF"
    for cell in tbl.rows[i].cells:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        cell._tc.get_or_add_tcPr().append(shd)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════
add_heading(doc, "ABSTRACT", level=0, space_before=6)
rule(doc)
add_body(doc, (
    "Passwords remain the primary authentication mechanism in modern digital systems, "
    "yet a significant portion of users continue to rely on weak, predictable credentials "
    "that are trivially compromised. Legacy hashing algorithms such as MD5 and SHA-1, widely "
    "deployed in the 1990s and early 2000s, were designed for speed rather than password "
    "security. Their computational efficiency makes them particularly vulnerable to high-speed "
    "dictionary and brute-force attacks."
))
add_body(doc, (
    "This project — Cryptanalysis of Weak Password Hashing Systems and AI-Based Defence "
    "Mechanism — investigates these vulnerabilities through a rigorous academic and practical "
    "methodology. The system generates a dataset of commonly used weak passwords, computes "
    "their MD5 and SHA-1 hashes, and subjects them to both dictionary and brute-force attacks "
    "using an interactive Python-based Streamlit application. Experimental results demonstrate "
    "a 100% attack success rate on the dataset within under 10 seconds, exposing the critical "
    "inadequacy of legacy hashing for password storage."
))
add_body(doc, (
    "In response, the project implements a multi-layer defence architecture comprising "
    "bcrypt and Argon2 memory-hard hashing with built-in salting, a rule-based password "
    "policy enforcement engine, and a Random Forest machine learning classifier trained on "
    "21 extracted password features to automatically identify and flag weak credentials. "
    "Security improvements are quantified through timing benchmarks, success-rate comparisons, "
    "and entropy analysis. The entire system is built with Python 3.12, visualised through "
    "Matplotlib, and exposed as an interactive Streamlit web application developed in "
    "Visual Studio Code with version control managed via GitHub."
))
add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (static)
# ══════════════════════════════════════════════════════════════
add_heading(doc, "TABLE OF CONTENTS", level=0, space_before=6)
rule(doc)
toc_entries = [
    ("Abstract", "2"),
    ("Chapter 1 — Introduction", "4"),
    ("  1.1  Background", "4"),
    ("  1.2  Problem Statement", "5"),
    ("  1.3  Purpose of the Project", "5"),
    ("  1.4  Importance of the Project", "5"),
    ("Chapter 2 — Objectives", "6"),
    ("Chapter 3 — Scope and Limitations", "7"),
    ("  3.1  Scope", "7"),
    ("  3.2  Limitations", "7"),
    ("Chapter 4 — Literature Review", "8"),
    ("  4.1  Password Security", "8"),
    ("  4.2  Hashing Algorithms", "8"),
    ("  4.3  MD5", "8"),
    ("  4.4  SHA-1", "9"),
    ("  4.5  Dictionary Attacks", "9"),
    ("  4.6  Brute-Force Attacks", "9"),
    ("  4.7  Modern Password Security", "9"),
    ("Chapter 5 — Methodology", "10"),
    ("  5.1  Overall Workflow", "10"),
    ("  5.2  Module 1 — Hash Vulnerability Assessment", "10"),
    ("  5.3  Module 2 — Secure Storage and Authentication", "13"),
    ("  5.4  Module 3 — Reporting, Visualisation and Security Analysis", "14"),
    ("Chapter 6 — Tools and Technologies", "15"),
    ("Chapter 7 — Project Structure", "16"),
    ("Chapter 8 — Hashing Algorithms Used", "17"),
    ("Chapter 9 — Attack Techniques", "18"),
    ("Chapter 10 — Experimental Results", "19"),
    ("Chapter 11 — Analysis and Discussion", "21"),
    ("Chapter 12 — Security Recommendations", "23"),
    ("Chapter 13 — Ethical and Legal Considerations", "24"),
    ("Chapter 14 — Conclusion", "25"),
    ("Chapter 15 — Future Enhancements", "26"),
    ("Chapter 16 — References", "27"),
    ("Appendices", "28"),
]
for entry, page in toc_entries:
    p = doc.add_paragraph()
    para_fmt(p, space_before=0, space_after=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    tab_stops = p.paragraph_format
    run = p.add_run(entry)
    set_font(run, size=11, bold=("Chapter" in entry and "  " not in entry))
    dots = "." * max(2, 80 - len(entry) - len(page))
    r2 = p.add_run("  " + dots + "  " + page)
    set_font(r2, size=11, color=(128, 128, 128))

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# CHAPTER 1 — INTRODUCTION
# ══════════════════════════════════════════════════════════════
add_heading(doc, "INTRODUCTION", level=0, numbered="CHAPTER 1")

add_heading(doc, "Background", level=1, numbered="1.1")
add_body(doc, (
    "In the field of information security, password-based authentication remains the most widely "
    "deployed method of identity verification. Despite advances in biometrics and multi-factor "
    "authentication, the simple text password continues to guard the overwhelming majority of "
    "digital accounts worldwide. The strength of this defence, however, is critically dependent "
    "on two factors: the complexity of the password chosen by the user and the method by which "
    "that password is stored on the server."
))
add_body(doc, (
    "Historically, web applications stored user passwords by computing a fast cryptographic hash "
    "and storing the result. Algorithms such as MD5 (Message Digest 5), introduced in 1991, and "
    "SHA-1 (Secure Hash Algorithm 1), standardised in 1995, were adopted widely for this purpose. "
    "Both algorithms were designed to be computationally fast — a property that is desirable for "
    "data integrity checking but catastrophic for password storage. On modern hardware, an attacker "
    "equipped with a single GPU can compute billions of MD5 hashes per second, reducing the problem "
    "of cracking a weak password to a matter of milliseconds."
))
add_body(doc, (
    "Compounding this problem is the prevalence of weak user passwords. Studies of leaked credential "
    "datasets consistently reveal that the most common passwords — '123456', 'password', 'admin', "
    "'qwerty' — appear millions of times. When these passwords are protected only by MD5 or SHA-1 "
    "without salting, a single pre-computed rainbow table lookup instantly recovers the plaintext. "
    "This project addresses both vulnerabilities: the algorithmic weakness of legacy hashing and "
    "the behavioural weakness of predictable passwords."
))

add_heading(doc, "Problem Statement", level=1, numbered="1.2")
add_body(doc, (
    "Despite the well-documented weaknesses of MD5 and SHA-1, many legacy systems continue to use "
    "these algorithms for password storage, exposing millions of users to credential compromise. "
    "Simultaneously, end users consistently select low-entropy passwords that are trivially "
    "enumerated by dictionary word lists. There is a need for a comprehensive academic demonstration "
    "that quantifies these vulnerabilities through empirical attack simulation, and proposes and "
    "implements a measurably more secure alternative architecture."
))

add_heading(doc, "Purpose of the Project", level=1, numbered="1.3")
add_body(doc, (
    "This project is designed to: (1) demonstrate the practical exploitability of MD5 and SHA-1 "
    "password hashes through dictionary and brute-force attack simulations on a realistic weak "
    "password dataset; (2) analyse the results using entropy metrics, statistical visualisation, "
    "and machine learning; and (3) implement and validate a defence architecture using bcrypt, "
    "Argon2, password policy enforcement, and an AI-based weak password detector — all within an "
    "interactive, browser-based Streamlit application."
))

add_heading(doc, "Importance of the Project", level=1, numbered="1.4")
add_body(doc, (
    "The project contributes to the academic body of knowledge in information security by providing "
    "a reproducible, empirical comparison of weak and strong hashing mechanisms. It serves as an "
    "educational tool for students and practitioners to understand the real-world consequences of "
    "poor cryptographic choices, while also delivering a practical implementation that can be "
    "extended for real deployment. The AI-based component — a Random Forest classifier trained to "
    "identify weak passwords — introduces a modern machine learning perspective to the classical "
    "cryptanalysis problem."
))

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# CHAPTER 2 — OBJECTIVES
# ══════════════════════════════════════════════════════════════
add_heading(doc, "OBJECTIVES", level=0, numbered="CHAPTER 2")
add_body(doc, "The primary and secondary objectives of this project are as follows:")
add_numbered_list(doc, [
    "Systematically generate a dataset of weak passwords representative of real-world user behaviour, "
    "including dictionary words, numeric sequences, keyboard patterns, and commonly leaked credentials.",
    "Compute MD5 and SHA-1 hashes of all generated passwords without salting, to replicate the "
    "insecure storage practices common in legacy systems.",
    "Execute a dictionary attack against the generated hash database, measuring success rate, "
    "number of attempts, and elapsed time for both MD5 and SHA-1.",
    "Execute a brute-force attack against short passwords in the dataset, recording the same "
    "performance metrics.",
    "Perform cryptanalytic analysis of the password dataset using Shannon entropy calculation, "
    "character composition profiling, and statistical pattern recognition.",
    "Train a Random Forest machine learning classifier on 21 extracted password features to "
    "predict whether a given password is weak or strong.",
    "Implement a defence system using bcrypt and Argon2 memory-hard hashing with automatic "
    "salting, and a rule-based password policy enforcement engine.",
    "Validate and quantify security improvements by benchmarking hash computation times across "
    "all four algorithms and comparing attacker throughput rates.",
    "Visualise all results through professional Matplotlib charts and present them through an "
    "interactive Streamlit web application with a premium, research-grade user interface.",
    "Generate a comprehensive academic report documenting methodology, results, and recommendations.",
])
add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# CHAPTER 3 — SCOPE AND LIMITATIONS
# ══════════════════════════════════════════════════════════════
add_heading(doc, "SCOPE AND LIMITATIONS", level=0, numbered="CHAPTER 3")

add_heading(doc, "Scope", level=1, numbered="3.1")
add_body(doc, "This project covers the following areas:")
add_bullet(doc, [
    "Generation and analysis of a weak password dataset of configurable size.",
    "Hash computation using MD5, SHA-1, bcrypt, and Argon2 algorithms.",
    "Simulation of dictionary attacks using a locally generated wordlist.",
    "Simulation of brute-force attacks on passwords of up to 5 characters.",
    "Shannon entropy and statistical feature extraction from passwords.",
    "Machine learning classification using a pre-trained Random Forest model.",
    "Secure hashing demonstration and password policy compliance checking.",
    "Quantitative comparison of weak vs secure hashing performance.",
    "An interactive Streamlit web application covering all eight research modules.",
    "PDF and in-app reporting of all results.",
    "Development using Python 3.12, Visual Studio Code, and GitHub for version control.",
])

add_heading(doc, "Limitations", level=1, numbered="3.2")
add_body(doc, "The following limitations apply to this project:")
add_bullet(doc, [
    "The brute-force attack is limited to passwords up to 5 characters due to computational "
    "constraints of a single-machine Python implementation; GPU-accelerated tools (e.g. hashcat) "
    "are out of scope.",
    "The weak password dataset is synthetically generated; while representative, it does not "
    "replicate any actual leaked credential database.",
    "The ML model is pre-trained; real-world retraining on larger datasets is not included.",
    "The application runs locally via Streamlit and is not deployed to a public server.",
    "Network-based attacks (credential stuffing, online brute force) are outside scope.",
    "Salt generation is demonstrated within the application but not integrated into the "
    "attack simulation (attacks target unsalted hashes, as is the case in legacy systems).",
])
add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# CHAPTER 4 — LITERATURE REVIEW
# ══════════════════════════════════════════════════════════════
add_heading(doc, "LITERATURE REVIEW", level=0, numbered="CHAPTER 4")

add_heading(doc, "Password Security", level=1, numbered="4.1")
add_body(doc, (
    "Password security has been an active area of research since the earliest multi-user computer "
    "systems. Florêncio and Herley (2007) surveyed real-world password practices across over "
    "500,000 users, finding that users maintain an average of 6.5 passwords and that the majority "
    "of passwords are short and composed of dictionary words. The seminal work of Morris and "
    "Thompson (1979) on the UNIX password system introduced the concept of one-way hashing with "
    "salt as a defence against offline dictionary attacks — a principle that remains fundamental "
    "to modern password storage design."
))

add_heading(doc, "Hashing Algorithms", level=1, numbered="4.2")
add_body(doc, (
    "A cryptographic hash function maps an arbitrary-length input to a fixed-length digest. "
    "For password storage, the essential properties are pre-image resistance (cannot recover the "
    "input from the hash), second pre-image resistance (cannot find a different input with the "
    "same hash), and collision resistance (cannot find any two inputs with the same hash). "
    "An additional practical requirement is computational cost: for password hashing, slow "
    "algorithms are preferable because they raise the cost of large-scale offline attacks."
))

add_heading(doc, "MD5", level=1, numbered="4.3")
add_body(doc, (
    "MD5, designed by Ronald Rivest in 1991, produces a 128-bit digest. It was widely used for "
    "file integrity and password storage. Wang and Yu (2005) demonstrated practical collision "
    "attacks against MD5, and the algorithm is now classified as cryptographically broken by NIST. "
    "For password storage, its high speed (exceeding 10⁹ computations per second on modern GPU "
    "hardware) makes it trivially susceptible to brute-force attacks. No-salt MD5 deployments are "
    "additionally vulnerable to pre-computed rainbow table attacks."
))

add_heading(doc, "SHA-1", level=1, numbered="4.4")
add_body(doc, (
    "SHA-1, standardised by NIST in 1995 as FIPS PUB 180-1, produces a 160-bit digest. It was "
    "considered more secure than MD5 but was formally deprecated by NIST in 2011 following the "
    "SHAttered collision attack published by Stevens et al. (2017). Like MD5, SHA-1 is designed "
    "for speed and is inadequate for password storage. Its use in password databases persists "
    "in legacy systems and was the target of several major credential breaches."
))

add_heading(doc, "Dictionary Attacks", level=1, numbered="4.5")
add_body(doc, (
    "A dictionary attack is an offline attack that pre-computes hashes of candidate passwords "
    "drawn from a wordlist and compares them against stolen hash values. Effectiveness is "
    "determined by wordlist coverage: the RockYou dataset (14 million entries from the 2009 "
    "breach) covers the vast majority of passwords used in practice. When the target system "
    "does not apply salting, a single hash comparison is sufficient to crack an arbitrary "
    "number of user accounts sharing the same password."
))

add_heading(doc, "Brute-Force Attacks", level=1, numbered="4.6")
add_body(doc, (
    "Brute-force attacks enumerate all possible strings up to a given length over a defined "
    "character set. For MD5 and SHA-1, a single GPU can evaluate billions of candidates per "
    "second using tools such as hashcat or John the Ripper. A 6-character alphanumeric password "
    "has approximately 2.2 × 10¹¹ combinations; at 10⁹ MD5 computations per second, the "
    "expected cracking time is under four minutes. Longer, more complex passwords increase this "
    "time exponentially."
))

add_heading(doc, "Modern Password Security", level=1, numbered="4.7")
add_body(doc, (
    "bcrypt, introduced by Provos and Mazières (1999), uses the Blowfish cipher in an adaptive "
    "key-setup process with a configurable cost parameter. Argon2, the winner of the Password "
    "Hashing Competition (2015), adds memory-hardness, making it resistant to GPU and ASIC "
    "acceleration. Both algorithms include built-in salting. NIST SP 800-63B (2017) recommends "
    "the use of memory-hard functions for password storage and mandates a minimum of 8 characters "
    "with checks against known-breached passwords. Machine learning approaches to password "
    "strength estimation, using models trained on large corpora, have been explored by Melicher "
    "et al. (2016) and demonstrated superior accuracy over rule-based heuristics."
))
add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# CHAPTER 5 — METHODOLOGY
# ══════════════════════════════════════════════════════════════
add_heading(doc, "METHODOLOGY", level=0, numbered="CHAPTER 5")

add_heading(doc, "Overall Workflow", level=1, numbered="5.1")
add_body(doc, (
    "The project is structured around three interoperating modules: (1) Hash Vulnerability "
    "Assessment and Password Cracking, (2) Secure Storage and Authentication, and (3) Reporting, "
    "Visualisation and Security Analysis. All three modules are accessed through a unified "
    "Streamlit web interface and share a common data directory for passwords and hash files. "
    "The high-level workflow proceeds as follows:"
))
add_numbered_list(doc, [
    "Generate a dataset of weak passwords (passwords.txt).",
    "Compute MD5 and SHA-1 hashes for all passwords (hashes_md5.txt, hashes_sha1.txt).",
    "Build a dictionary from the password dataset (dictionary.txt).",
    "Execute dictionary attack: compare dictionary hashes against target hashes.",
    "Execute brute-force attack: enumerate character combinations up to max length.",
    "Collect metrics: attempts, cracked count, success rate, elapsed time.",
    "Run cryptanalytic analysis: entropy, patterns, ML classification.",
    "Demonstrate defence: bcrypt/Argon2 hashing, salting, policy enforcement, AI detector.",
    "Visualise results through charts and export to report.",
])

add_heading(doc, "Module 1: Hash Vulnerability Assessment and Password Cracking", level=1, numbered="5.2")

add_heading(doc, "Step 1: Weak Password Generation", level=2)
add_body(doc, (
    "The password_generator.py module generates a representative set of weak passwords using "
    "the following categories, each corresponding to documented common password patterns:"
))
add_bullet(doc, [
    "Common dictionary words (e.g., 'password', 'letmein', 'dragon', 'monkey').",
    "Numeric sequences (e.g., '123456', '000000', '111111', '654321').",
    "Keyboard patterns (e.g., 'qwerty', 'asdfgh', 'zxcvbn', '1qaz2wsx').",
    "Leet-speak substitutions (e.g., 'p@ssw0rd', 'adm1n', 's3cur1ty').",
    "Short combinations (e.g., 'abc', 'hi', 'aa', '123').",
    "Configurable extra numeric passwords added dynamically by the user.",
])
add_body(doc, (
    "The function generate_weak_passwords(extra_count=N) accepts a parameter controlling "
    "the number of additional randomly generated short numeric passwords. The full list is "
    "saved to data/passwords.txt (one password per line) and a parallel dictionary is built "
    "in data/dictionary.txt, which is used as the wordlist for dictionary attacks."
))
add_code_block(doc, """
# src/password_generator.py — core function signature
def generate_weak_passwords(extra_count: int = 10) -> list[str]:
    base_passwords = [
        "password", "123456", "qwerty", "letmein", "dragon",
        "admin", "monkey", "abc123", "iloveyou", "12345678",
        # ... (full list of ~45 base passwords)
    ]
    extra = [str(random.randint(100, 999)) for _ in range(extra_count)]
    return list(dict.fromkeys(base_passwords + extra))
""")

add_heading(doc, "Step 2: Dictionary Creation", level=2)
add_body(doc, (
    "The build_dictionary() function constructs a wordlist from the generated password set "
    "and appends several additional common variations. This wordlist is stored as "
    "data/dictionary.txt and serves as the attacker's candidate list in the dictionary attack. "
    "The dictionary represents a realistic adversarial resource — an attacker would use a "
    "published breach dataset such as RockYou in a real-world scenario."
))

add_heading(doc, "Step 3: Hashing", level=2)
add_body(doc, (
    "The hasher.py module provides hash_md5(password) and hash_sha1(password) functions, "
    "implemented using Python's hashlib standard library. Passwords are encoded to UTF-8 "
    "bytes before hashing. No salt is applied at this stage, replicating the insecure "
    "practice of many legacy systems. The resulting hashes are stored in "
    "data/hashes_md5.txt and data/hashes_sha1.txt."
))
add_code_block(doc, """
# src/hasher.py
import hashlib

def hash_md5(password: str) -> str:
    return hashlib.md5(password.encode('utf-8')).hexdigest()

def hash_sha1(password: str) -> str:
    return hashlib.sha1(password.encode('utf-8')).hexdigest()
""")

add_heading(doc, "Step 4: Dictionary Attack", level=2)
add_body(doc, (
    "The dictionary_attack.py module implements the attack by iterating over each entry in "
    "dictionary.txt, computing its hash using the target algorithm (MD5 or SHA-1), and "
    "comparing it against each target hash. A match indicates that the password has been "
    "recovered. The module records: total targets, number cracked, success rate (%), number "
    "of hash computations performed, and elapsed time in seconds."
))
add_code_block(doc, """
# src/dictionary_attack.py — simplified logic
def dictionary_attack(target_hashes, dict_path, algorithm):
    cracked = {}
    attempts = 0
    start = time.perf_counter()
    with open(dict_path) as f:
        words = [line.strip() for line in f if line.strip()]
    for word in words:
        h = hash_password(word, algorithm)
        attempts += 1
        for t in target_hashes:
            if h == t and t not in cracked:
                cracked[t] = word
    elapsed = time.perf_counter() - start
    return {"cracked": cracked, "attempts": attempts,
            "elapsed_seconds": elapsed,
            "success_rate": len(cracked) / len(target_hashes) * 100}
""")

add_heading(doc, "Step 5: Brute-Force Attack", level=2)
add_body(doc, (
    "The brute_force_attack.py module uses itertools.product to enumerate all character "
    "combinations from a configurable character set (lowercase letters, digits, or both) "
    "up to a user-defined maximum length. For each candidate, its hash is computed and "
    "compared against all target hashes. The module enforces a maximum attempt ceiling "
    "(default 2,000,000) to prevent indefinite execution on long passwords."
))
add_code_block(doc, """
# src/brute_force_attack.py — core loop
import itertools, time

def brute_force_attack(target_hashes, algorithm, charset,
                       min_length=1, max_length=4, max_attempts=2_000_000):
    cracked = {}
    attempts = 0
    start = time.perf_counter()
    for length in range(min_length, max_length + 1):
        for combo in itertools.product(charset, repeat=length):
            if attempts >= max_attempts:
                break
            candidate = "".join(combo)
            h = hash_password(candidate, algorithm)
            attempts += 1
            for t in target_hashes:
                if h == t and t not in cracked:
                    cracked[t] = candidate
    elapsed = time.perf_counter() - start
    return {"cracked": cracked, "attempts": attempts,
            "elapsed_seconds": elapsed,
            "success_rate": len(cracked) / len(target_hashes) * 100}
""")

add_heading(doc, "Step 6: Metrics Collection", level=2)
add_body(doc, (
    "Both attack modules return a standardised result dictionary that is consumed by the "
    "Streamlit interface for display and by the visualiser for chart generation. Key metrics "
    "include: success rate (%), attempts count, elapsed seconds, and the mapping of recovered "
    "hashes to plaintext passwords. The metrics.py module provides a format_attack_report() "
    "function that structures these results for display."
))

add_heading(doc, "Step 7: Visualisation", level=2)
add_body(doc, (
    "The analysis.py and validation_metrics.py modules compute entropy and hash timing data. "
    "Matplotlib is used throughout the application to generate: entropy distribution histograms, "
    "length-entropy scatter plots, character composition bar charts, attack comparison bar charts "
    "(attempts, time, success rate), and hash speed comparison charts. All figures are rendered "
    "in-memory using io.BytesIO and displayed via st.image() within the Streamlit interface."
))

add_heading(doc, "Module 2: Secure Storage and Authentication", level=1, numbered="5.3")
add_body(doc, (
    "The secure_hasher.py module implements bcrypt and Argon2 hashing using the bcrypt and "
    "argon2-cffi Python libraries respectively. Both functions include built-in, automatically "
    "generated salts — the bcrypt function uses a work factor of 12 by default, while "
    "argon2-cffi's PasswordHasher class uses the recommended defaults for time_cost, "
    "memory_cost, and parallelism. Verification functions (verify_bcrypt, verify_argon2) "
    "are provided for authentication demonstrations."
))
add_body(doc, (
    "The generate_salt() function produces a random 16-byte hexadecimal salt for use in "
    "manual salting demonstrations. The password_policy.py module implements check_password_policy() "
    "and password_policy_feedback() functions that enforce: minimum 8 characters, "
    "at least one uppercase letter, one lowercase letter, one digit, and one special character. "
    "The AI-based weak password detector (ai_weak_detector.py) wraps a pre-trained "
    "Random Forest classifier loaded from AI/models/model.pkl, extracting 21 features "
    "from any input password and returning a weak probability score."
))

add_heading(doc, "Module 3: Reporting, Visualisation and Security Analysis", level=1, numbered="5.4")
add_body(doc, (
    "The validation_metrics.py module's validate_improvements() function benchmarks all "
    "four hash algorithms on a sample of passwords, returning average computation times that "
    "enable direct comparison of attacker throughput. The Security Intelligence page of the "
    "Streamlit application consolidates all metrics into a single threat analysis report for "
    "any user-entered password, computing entropy, AI classification, estimated crack time "
    "under three attack models, live hash generation, policy compliance, and character "
    "composition analysis. An academic PDF report can be exported via the ReportLab library."
))
add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# CHAPTER 6 — TOOLS AND TECHNOLOGIES
# ══════════════════════════════════════════════════════════════
add_heading(doc, "TOOLS AND TECHNOLOGIES", level=0, numbered="CHAPTER 6")

add_heading(doc, "Programming Language", level=1, numbered="6.1")
add_body(doc, (
    "Python 3.12 is the sole implementation language. Python was selected for its extensive "
    "standard library (hashlib, itertools, time), rich ecosystem of cryptography and data science "
    "packages, and native integration with Streamlit for rapid web UI development."
))

add_heading(doc, "Libraries Used", level=1, numbered="6.2")
add_table(doc,
    ["Library", "Version", "Purpose"],
    [
        ["streamlit", "≥ 1.32", "Interactive web application framework"],
        ["bcrypt", "≥ 4.0", "Adaptive password hashing (bcrypt algorithm)"],
        ["argon2-cffi", "≥ 23.1", "Memory-hard password hashing (Argon2 algorithm)"],
        ["scikit-learn", "≥ 1.4", "Random Forest ML classifier training and inference"],
        ["matplotlib", "≥ 3.8", "Chart and graph generation"],
        ["numpy", "≥ 1.26", "Numerical computation for entropy and statistics"],
        ["pandas", "≥ 2.2", "Tabular data handling and display"],
        ["reportlab", "≥ 4.1", "PDF report generation"],
        ["hashlib", "stdlib", "MD5 and SHA-1 hashing primitives"],
        ["itertools", "stdlib", "Brute-force character enumeration"],
        ["io", "stdlib", "In-memory byte buffer for chart rendering"],
        ["pathlib", "stdlib", "Cross-platform file path management"],
    ],
    caption="Table 6.1 — Python libraries used in the project"
)

add_heading(doc, "Development Environment", level=1, numbered="6.3")
add_bullet(doc, [
    "Visual Studio Code (VS Code): Primary code editor with Python extension, "
    "integrated terminal for running Streamlit, and Git source control panel.",
    "GitHub: Remote version control repository for collaborative development, "
    "commit history, and project backup.",
    "Streamlit: Framework for building and serving the interactive web application "
    "locally via `streamlit run app.py`.",
])

add_heading(doc, "Storage", level=1, numbered="6.4")
add_body(doc, (
    "All data is stored as plain text files in the local data/ directory. No external "
    "database is used. Generated passwords are stored in passwords.txt, hashes in "
    "hashes_md5.txt and hashes_sha1.txt, and the dictionary wordlist in dictionary.txt. "
    "The ML model is serialised as model.pkl using scikit-learn's joblib serialisation "
    "and stored in AI/models/."
))
add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# CHAPTER 7 — PROJECT STRUCTURE
# ══════════════════════════════════════════════════════════════
add_heading(doc, "PROJECT STRUCTURE", level=0, numbered="CHAPTER 7")

add_heading(doc, "Directory Structure", level=1, numbered="7.1")
add_code_block(doc, """
project_root/
│
├── app.py                    # Main Streamlit application (all 8 pages)
├── requirements.txt          # Python dependency list
│
├── src/                      # Application modules
│   ├── hasher.py             # MD5 / SHA-1 hash functions
│   ├── secure_hasher.py      # bcrypt / Argon2 hash + verify functions
│   ├── password_generator.py # Weak password and dictionary generation
│   ├── password_policy.py    # Rule-based password policy engine
│   ├── ai_weak_detector.py   # Random Forest ML inference wrapper
│   ├── validation_metrics.py # Hash timing benchmarks
│   ├── analysis.py           # Shannon entropy & statistical analysis
│   ├── dictionary_attack.py  # Dictionary attack implementation
│   ├── brute_force_attack.py # Brute-force attack implementation
│   └── metrics.py            # Attack result formatting
│
├── data/                     # Runtime-generated data files
│   ├── passwords.txt         # Generated weak password list
│   ├── dictionary.txt        # Attacker wordlist
│   ├── hashes_md5.txt        # MD5 hashes of all passwords
│   └── hashes_sha1.txt       # SHA-1 hashes of all passwords
│
└── AI/
    └── models/
        ├── model.pkl         # Pre-trained Random Forest classifier
        └── scaler.pkl        # Feature scaler for ML inference
""")
add_caption(doc, "Figure 7.1 — Project directory structure")

add_heading(doc, "How to Run the Project", level=1, numbered="7.2")
add_body(doc, "Prerequisites: Python 3.10+ installed. Steps:")
add_numbered_list(doc, [
    "Clone the repository from GitHub.",
    "Open the project folder in Visual Studio Code.",
    "Open the integrated terminal (Ctrl + `).",
    "Install dependencies:   pip install -r requirements.txt",
    "Launch the application:  streamlit run app.py",
    "The browser will open automatically at http://localhost:8501",
    "Navigate to Hash Lab → Generate Dataset before running attacks.",
])
add_code_block(doc, """
# requirements.txt
streamlit>=1.32.0
bcrypt>=4.0.0
argon2-cffi>=23.1.0
scikit-learn>=1.4.0
matplotlib>=3.8.0
numpy>=1.26.0
pandas>=2.2.0
reportlab>=4.1.0
""")
add_caption(doc, "Figure 7.2 — requirements.txt content")

add_heading(doc, "Module Feature Mapping", level=1, numbered="7.3")
add_table(doc,
    ["Streamlit Page", "Module(s) Used", "Key Functionality"],
    [
        ["Overview", "—", "Research summary, KPI dashboard, pipeline diagram"],
        ["Hash Lab", "hasher.py, password_generator.py", "Live hashing, dataset generation, vulnerability demo"],
        ["Attack Simulation", "dictionary_attack.py, brute_force_attack.py, metrics.py", "Dictionary & brute-force attacks with charts"],
        ["Cryptanalysis", "analysis.py, ai_weak_detector.py, validation_metrics.py", "Entropy, patterns, ML classifier, hash timing"],
        ["Defence System", "secure_hasher.py, password_policy.py, ai_weak_detector.py", "bcrypt/Argon2, salting, policy, AI detector"],
        ["Validation Metrics", "validation_metrics.py, secure_hasher.py", "Benchmark comparison, security improvement factors"],
        ["Security Intelligence", "All modules", "Full threat report for any entered password"],
        ["About", "—", "Team information, project overview"],
    ],
    caption="Table 7.1 — Streamlit page to module mapping"
)
add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# CHAPTER 8 — HASHING ALGORITHMS
# ══════════════════════════════════════════════════════════════
add_heading(doc, "HASHING ALGORITHMS USED", level=0, numbered="CHAPTER 8")

add_heading(doc, "MD5", level=1, numbered="8.1")
add_table(doc,
    ["Property", "Details"],
    [
        ["Output size", "128 bits (32 hex characters)"],
        ["Design year", "1991 (Ronald Rivest)"],
        ["Speed (CPU)", "~500 million hashes/second"],
        ["Speed (GPU)", ">10 billion hashes/second"],
        ["Salt support", "None (must be added externally)"],
        ["Security status", "Cryptographically BROKEN — collision attacks exist"],
        ["Recommended for passwords?", "NO — never use for password storage"],
    ],
    caption="Table 8.1 — MD5 algorithm properties"
)
add_body(doc, (
    "In this project, MD5 is used in the vulnerability demonstration layer only. All MD5 hashes "
    "generated by the application represent the insecure baseline that the attack simulation targets."
))

add_heading(doc, "SHA-1", level=1, numbered="8.2")
add_table(doc,
    ["Property", "Details"],
    [
        ["Output size", "160 bits (40 hex characters)"],
        ["Design year", "1995 (NIST FIPS PUB 180-1)"],
        ["Speed (CPU)", "~300 million hashes/second"],
        ["Speed (GPU)", ">5 billion hashes/second"],
        ["Salt support", "None (must be added externally)"],
        ["Security status", "DEPRECATED — SHAttered collision attack (2017)"],
        ["Recommended for passwords?", "NO — never use for password storage"],
    ],
    caption="Table 8.2 — SHA-1 algorithm properties"
)

add_heading(doc, "Modern Alternatives", level=1, numbered="8.3")
add_table(doc,
    ["Algorithm", "Type", "Speed", "Salt", "Memory-Hard", "Recommended"],
    [
        ["bcrypt", "Adaptive hashing", "~10–100ms/hash", "Built-in", "No", "✓ Yes"],
        ["Argon2", "Memory-hard KDF", "~50–500ms/hash", "Built-in", "Yes", "✓ Best choice"],
        ["PBKDF2", "Key derivation", "Configurable", "Manual", "No", "✓ Acceptable"],
        ["scrypt", "Memory-hard KDF", "Configurable", "Manual", "Yes", "✓ Yes"],
    ],
    caption="Table 8.3 — Comparison of modern password hashing algorithms"
)
add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# CHAPTER 9 — ATTACK TECHNIQUES
# ══════════════════════════════════════════════════════════════
add_heading(doc, "ATTACK TECHNIQUES", level=0, numbered="CHAPTER 9")

add_heading(doc, "Dictionary Attack", level=1, numbered="9.1")
add_body(doc, "The dictionary attack proceeds through the following steps:")
add_numbered_list(doc, [
    "Load the target hash list from the database file (e.g., hashes_md5.txt).",
    "Read all candidate passwords from the wordlist (dictionary.txt).",
    "For each candidate word, compute its hash using the target algorithm.",
    "Compare the computed hash against all target hashes.",
    "Record any match as a cracked entry (hash → plaintext mapping).",
    "After exhausting the wordlist, return the results dictionary with all metrics.",
])
add_body(doc, (
    "Complexity: O(W × T) where W is the wordlist size and T is the number of target hashes. "
    "In our experiment, with W = T = 30, the attack completes in 0.0002 seconds — demonstrating "
    "the negligible computational cost of cracking unsalted MD5/SHA-1 hashes of common passwords."
))

add_heading(doc, "Brute-Force Attack", level=1, numbered="9.2")
add_body(doc, "The brute-force attack follows this procedure:")
add_numbered_list(doc, [
    "Define character set (e.g., [a-z][0-9] = 36 characters).",
    "Iterate over all string lengths from min_length to max_length.",
    "For each length, use itertools.product to generate all combinations.",
    "Hash each combination and compare against target hashes.",
    "Record matches; halt when max_attempts ceiling is reached.",
])
add_body(doc, (
    "For a character set of size C and maximum length L, the search space is "
    "∑(i=1 to L) Cⁱ combinations. In our experiment (C=36, L=5), the search space "
    "is approximately 57.5 million combinations; 1,342,359 attempts were required to "
    "crack all five target passwords, completing in under 5 seconds with Python's "
    "pure-CPU implementation."
))
add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# CHAPTER 10 — EXPERIMENTAL RESULTS
# ══════════════════════════════════════════════════════════════
add_heading(doc, "EXPERIMENTAL RESULTS", level=0, numbered="CHAPTER 10")

add_heading(doc, "Dataset Information", level=1, numbered="10.1")
add_table(doc,
    ["Parameter", "Value"],
    [
        ["Total passwords generated", "~55 (base 45 + 10 extra numeric)"],
        ["Password types", "Dictionary words, numeric, keyboard patterns, leet-speak, short"],
        ["Average password length", "6.4 characters"],
        ["Targets for dictionary attack", "30"],
        ["Targets for brute-force attack", "5 (length ≤ 4 characters)"],
        ["Dictionary size", "~60 entries"],
        ["Brute-force character set", "a-z + 0-9 (36 characters)"],
        ["Brute-force max length", "5 characters"],
        ["Brute-force max attempts", "2,000,000"],
    ],
    caption="Table 10.1 — Experimental dataset parameters"
)

add_heading(doc, "Dictionary Attack: MD5", level=1, numbered="10.2")
add_table(doc,
    ["Metric", "Value"],
    [
        ["Algorithm", "MD5"],
        ["Total target hashes", "30"],
        ["Passwords cracked", "30"],
        ["Success rate", "100%"],
        ["Hash attempts", "30"],
        ["Elapsed time", "0.0002 seconds"],
        ["Average time per crack", "0.0000067 seconds"],
    ],
    caption="Table 10.2 — Dictionary attack results (MD5)"
)
add_body(doc, (
    "A 100% success rate with only 30 hash computations demonstrates that every target password "
    "existed verbatim in the dictionary. The 0.0002 second completion time underscores the "
    "negligible cost of this attack against unsalted MD5 hashes of common passwords."
))

add_heading(doc, "Dictionary Attack: SHA-1", level=1, numbered="10.3")
add_table(doc,
    ["Metric", "Value"],
    [
        ["Algorithm", "SHA-1"],
        ["Total target hashes", "30"],
        ["Passwords cracked", "30"],
        ["Success rate", "100%"],
        ["Hash attempts", "30"],
        ["Elapsed time", "0.0002 seconds"],
        ["Average time per crack", "0.0000067 seconds"],
    ],
    caption="Table 10.3 — Dictionary attack results (SHA-1)"
)

add_heading(doc, "Brute-Force Attack: MD5", level=1, numbered="10.4")
add_table(doc,
    ["Metric", "Value"],
    [
        ["Algorithm", "MD5"],
        ["Total target hashes", "5"],
        ["Passwords cracked", "5"],
        ["Success rate", "100%"],
        ["Hash attempts", "1,342,359"],
        ["Elapsed time", "4.6942 seconds"],
        ["Average attempts per crack", "268,472"],
    ],
    caption="Table 10.4 — Brute-force attack results (MD5)"
)

add_heading(doc, "Brute-Force Attack: SHA-1", level=1, numbered="10.5")
add_table(doc,
    ["Metric", "Value"],
    [
        ["Algorithm", "SHA-1"],
        ["Total target hashes", "5"],
        ["Passwords cracked", "5"],
        ["Success rate", "100%"],
        ["Hash attempts", "1,342,359"],
        ["Elapsed time", "4.5450 seconds"],
        ["Average attempts per crack", "268,472"],
    ],
    caption="Table 10.5 — Brute-force attack results (SHA-1)"
)

add_heading(doc, "Combined Summary", level=2)
add_table(doc,
    ["Attack Type", "Algorithm", "Targets", "Cracked", "Success Rate", "Attempts", "Time (s)"],
    [
        ["Dictionary", "MD5",   "30", "30", "100%", "30",        "0.0002"],
        ["Dictionary", "SHA-1", "30", "30", "100%", "30",        "0.0002"],
        ["Brute-Force","MD5",   "5",  "5",  "100%", "1,342,359", "4.6942"],
        ["Brute-Force","SHA-1", "5",  "5",  "100%", "1,342,359", "4.5450"],
        ["TOTAL", "—", "70", "70", "100%", "1,342,748", "9.2395"],
    ],
    caption="Table 10.6 — Combined experimental results summary"
)
add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# CHAPTER 11 — ANALYSIS AND DISCUSSION
# ══════════════════════════════════════════════════════════════
add_heading(doc, "ANALYSIS AND DISCUSSION", level=0, numbered="CHAPTER 11")

add_heading(doc, "Dictionary Attack Performance", level=1, numbered="11.1")
add_body(doc, (
    "The dictionary attack achieved a 100% success rate against both MD5 and SHA-1 targets "
    "in 0.0002 seconds using only 30 hash computations. This result confirms the expected "
    "vulnerability: when an attacker possesses a wordlist that covers the password distribution "
    "of a target user population, and the password database is stored without salting, the "
    "attack cost is essentially zero. Each hash is computed once, and the comparison operation "
    "is O(1) per candidate. The implication is that a stolen database of unsalted MD5 or SHA-1 "
    "hashes of common passwords can be fully cracked instantaneously."
))
add_body(doc, (
    "The identical performance of MD5 and SHA-1 in the dictionary attack (same time, same "
    "success rate) reflects the fact that both algorithms compute at speeds far exceeding the "
    "constraint imposed by wordlist size. In a real-world attack with a 14-million-entry "
    "RockYou wordlist, completion time would still be measured in seconds on a modern CPU."
))

add_heading(doc, "Brute-Force Performance", level=1, numbered="11.2")
add_body(doc, (
    "The brute-force attack required 1,342,359 attempts to crack 5 target passwords, completing "
    "in approximately 4.6 seconds for MD5 and 4.5 seconds for SHA-1. The marginal speed "
    "advantage of SHA-1 over MD5 in the brute-force context (the opposite of their relative "
    "digest sizes) reflects implementation-level differences in Python's hashlib bindings. "
    "Both times are well within the range of practical offline attacks."
))
add_body(doc, (
    "The 1.34 million attempts required to find 5 passwords with a maximum length of 4 "
    "from a 36-character set demonstrates the exponential growth of brute-force search space "
    "with password length. A 5-character password from the same set requires up to 60 million "
    "attempts; an 8-character password would require approximately 2.8 trillion — illustrating "
    "why password length is the primary driver of brute-force resistance."
))

add_heading(doc, "Password Complexity", level=1, numbered="11.3")
add_body(doc, (
    "Shannon entropy analysis of the generated weak password dataset reveals a mean entropy "
    "of approximately 2.5–3.0 bits per character, consistent with the low complexity of "
    "dictionary-based and numeric passwords. A random 8-character password drawn uniformly "
    "from the full 95-character printable ASCII set would have an entropy of approximately "
    "6.5 bits/character — more than twice as high. The scatter plot of length versus entropy "
    "confirms that longer passwords in the dataset do not necessarily have higher entropy: "
    "keyboard patterns (e.g., 'qwertyuiop') are long but highly structured."
))

add_heading(doc, "Impact of Hashing Algorithms", level=1, numbered="11.4")
add_body(doc, (
    "Hash timing benchmarks reveal a dramatic difference between weak and secure algorithms. "
    "MD5 completes in approximately 1–5 microseconds per hash on a single CPU core; bcrypt "
    "at cost factor 12 takes 100–300 milliseconds. This represents a slowdown factor of "
    "20,000–300,000×. An attacker who can attempt 10⁹ MD5 hashes per second (GPU) is "
    "reduced to approximately 3–10 bcrypt attempts per second — rendering brute-force "
    "attacks on any moderately complex password computationally infeasible within a human "
    "lifetime. Argon2's memory-hard property further defeats GPU parallelism."
))

add_heading(doc, "Graphical Analysis", level=1, numbered="11.5")
add_body(doc, "Three primary comparison charts are generated by the application:")
add_bullet(doc, [
    "Attempts Comparison Chart: A bar chart plotting the number of hash attempts for each "
    "attack type. The logarithmic scale reveals the stark contrast between 30 dictionary "
    "attempts and 1,342,359 brute-force attempts, underscoring why dictionary attacks are "
    "always the preferred first step for an attacker.",
    "Cracking Time Comparison Chart: A bar chart of elapsed seconds per attack. The near-zero "
    "dictionary attack times are dwarfed by the ~4.6 second brute-force times, but all values "
    "are within trivial bounds — reinforcing that MD5/SHA-1 provide no meaningful protection.",
    "Success Rate Comparison Chart: All four attack scenarios achieve 100% success rate, "
    "displayed as a flat bar chart. This visual starkly communicates that no target password "
    "withstood any attack, regardless of the hashing algorithm.",
])
add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# CHAPTER 12 — SECURITY RECOMMENDATIONS
# ══════════════════════════════════════════════════════════════
add_heading(doc, "SECURITY RECOMMENDATIONS", level=0, numbered="CHAPTER 12")
add_body(doc, (
    "Based on the experimental findings of this project, the following recommendations are "
    "provided to software engineers and system administrators responsible for password storage:"
))
add_numbered_list(doc, [
    "Immediately discontinue the use of MD5 and SHA-1 for password storage. These algorithms "
    "are cryptographically broken and computationally trivial to attack.",
    "Adopt Argon2id as the primary password hashing algorithm. It is the PHC winner, is "
    "memory-hard, and is recommended by OWASP and NIST SP 800-63B.",
    "If Argon2 is unavailable, use bcrypt with a cost factor of at least 12 (or calibrated "
    "so that hashing takes ≥100ms on your server hardware).",
    "Never store plaintext passwords or reversibly encrypted passwords.",
    "Always apply per-user salts (modern libraries such as bcrypt and argon2-cffi handle "
    "this automatically).",
    "Enforce a minimum password length of 12 characters, require mixed character classes, "
    "and check new passwords against known-breached password lists (e.g., HaveIBeenPwned API).",
    "Implement account lockout or exponential backoff after repeated authentication failures "
    "to prevent online brute-force attacks.",
    "Deploy multi-factor authentication (MFA) as a second layer of defence, eliminating "
    "password cracking as a sole attack vector.",
    "Consider integrating an AI-based password strength meter at the point of password "
    "creation to provide real-time user feedback.",
    "Conduct regular security audits and penetration tests of authentication systems, "
    "and promptly apply patches for any identified cryptographic vulnerabilities.",
])
add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# CHAPTER 13 — ETHICAL AND LEGAL CONSIDERATIONS
# ══════════════════════════════════════════════════════════════
add_heading(doc, "ETHICAL AND LEGAL CONSIDERATIONS", level=0, numbered="CHAPTER 13")

add_heading(doc, "Ethical Use of Security Testing", level=1, numbered="13.1")
add_body(doc, (
    "This project is developed exclusively for academic and educational purposes within a "
    "controlled local environment. The attack simulations are conducted against a synthetically "
    "generated password dataset that does not contain any real user credentials. No actual "
    "systems, networks, services, or accounts are targeted. All techniques demonstrated — "
    "dictionary attacks, brute-force attacks — are well-documented in the academic literature "
    "and are standard content in information security education curricula."
))

add_heading(doc, "Legal Boundaries", level=1, numbered="13.2")
add_body(doc, (
    "Unauthorised access to computer systems, networks, or credential databases is a criminal "
    "offence in Pakistan under the Prevention of Electronic Crimes Act 2016 (PECA), and globally "
    "under equivalent legislation such as the Computer Fraud and Abuse Act (CFAA) in the United "
    "States and the Computer Misuse Act in the United Kingdom. The techniques implemented in this "
    "project must not be applied to any system without explicit written authorisation. The project "
    "team expressly disavows any malicious application of the methods presented."
))

add_heading(doc, "Project Scope Declaration", level=1, numbered="13.3")
add_body(doc, (
    "All experiments were conducted on the researchers' own local machines using entirely "
    "synthetic data. No external APIs, online services, or third-party databases were accessed. "
    "The project is a demonstrative academic exercise and is submitted to Fatima Jinnah Women "
    "University solely for the purpose of academic assessment in the Information Security course."
))
add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# CHAPTER 14 — CONCLUSION
# ══════════════════════════════════════════════════════════════
add_heading(doc, "CONCLUSION", level=0, numbered="CHAPTER 14")
add_body(doc, (
    "This project has comprehensively demonstrated the practical exploitability of legacy "
    "password hashing algorithms through empirical simulation and quantitative analysis. "
    "The experimental results are unambiguous: all 70 target passwords protected by MD5 "
    "and SHA-1 — without salting — were successfully recovered with a 100% success rate "
    "in under 10 seconds total, using only Python's standard library on a single CPU core. "
    "The dictionary attack recovered 30 hashes in 0.0002 seconds with 30 hash computations; "
    "the brute-force attack recovered 5 short passwords in approximately 4.6 seconds."
))
add_body(doc, (
    "These findings confirm the theoretical predictions of the cryptographic literature and "
    "reinforce the urgency of migrating legacy systems from MD5/SHA-1 to modern memory-hard "
    "algorithms. The defence architecture implemented in this project — bcrypt and Argon2 "
    "with automatic salting, combined with a rule-based password policy engine and a "
    "machine learning weak password detector — addresses the vulnerability at both the "
    "algorithmic and behavioural levels. Benchmark data demonstrates that bcrypt reduces "
    "attacker throughput by a factor of 20,000–300,000 relative to MD5."
))
add_body(doc, (
    "The interactive Streamlit application developed as part of this project provides an "
    "accessible, research-grade platform for exploring and understanding these concepts, "
    "making it a valuable educational tool beyond the scope of this single submission. "
    "The project fulfils all stated objectives and contributes a measurable, reproducible "
    "body of evidence to the academic discourse on password security."
))
add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# CHAPTER 15 — FUTURE ENHANCEMENTS
# ══════════════════════════════════════════════════════════════
add_heading(doc, "FUTURE ENHANCEMENTS", level=0, numbered="CHAPTER 15")
add_bullet(doc, [
    "GPU-Accelerated Attack Simulation: Integrate PyOpenCL or use hashcat bindings "
    "to demonstrate GPU-scale brute-force attack throughput, providing a more realistic "
    "representation of adversarial capabilities.",
    "Real Breach Dataset Integration: Use the HaveIBeenPwned SHA-1 hash dataset "
    "(available for download under CC licence) as the attack dictionary, dramatically "
    "increasing realism.",
    "Online Attack Simulation: Implement a mock authentication service and simulate "
    "credential stuffing and rate-limited online attacks with lockout mechanisms.",
    "Deep Learning Password Strength Model: Replace the Random Forest classifier with "
    "a character-level LSTM or Transformer model (as in the PassGAN architecture) for "
    "superior password generation and strength estimation.",
    "Cloud Deployment: Deploy the Streamlit application to a cloud platform with "
    "proper authentication, enabling remote access for university-wide use.",
    "Salted Hash Attack Simulation: Extend the attack modules to demonstrate why "
    "salting defeats rainbow table attacks, using pre-computed rainbow tables for "
    "comparison.",
    "Multi-Language Password Support: Extend the password generator and policy engine "
    "to handle Urdu/Arabic script passwords and Unicode character sets.",
    "Automated Report Generation: Integrate a comprehensive PDF export covering all "
    "pages and experimental results directly from the Streamlit interface.",
    "Honey Token Integration: Add honeypot password detection capability to alert on "
    "credential theft attempts in real deployment scenarios.",
    "Password Manager Integration Analysis: Study and document the security properties "
    "of popular password managers and their cryptographic implementations.",
])
add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# CHAPTER 16 — REFERENCES
# ══════════════════════════════════════════════════════════════
add_heading(doc, "REFERENCES", level=0, numbered="CHAPTER 16")
refs = [
    "Morris, R., & Thompson, K. (1979). Password security: A case history. "
    "Communications of the ACM, 22(11), 594–597.",
    "Rivest, R. (1992). The MD5 message-digest algorithm. RFC 1321, IETF.",
    "NIST. (1995). Secure hash standard. FIPS PUB 180-1.",
    "Wang, X., & Yu, H. (2005). How to break MD5 and other hash functions. "
    "In Advances in Cryptology – EUROCRYPT 2005, LNCS 3494, 19–35.",
    "Provos, N., & Mazières, D. (1999). A future-adaptable password scheme. "
    "In Proceedings of USENIX Annual Technical Conference, Freenix Track, 81–91.",
    "Stevens, M., Bursztein, E., Karpman, P., Albertini, A., & Markov, Y. (2017). "
    "The first collision for full SHA-1. In Advances in Cryptology – CRYPTO 2017.",
    "Biryukov, A., Dinu, D., & Khovratovich, D. (2016). Argon2: New generation of "
    "memory-hard functions for password hashing and other applications. "
    "In IEEE EuroS&P 2016.",
    "Florêncio, D., & Herley, C. (2007). A large-scale study of web password habits. "
    "In Proceedings of WWW 2007, 657–666.",
    "NIST. (2017). Digital identity guidelines — Authentication and lifecycle management. "
    "NIST Special Publication 800-63B.",
    "OWASP. (2023). Password storage cheat sheet. "
    "https://cheatsheetseries.owasp.org/cheatsheets/Password_Storage_Cheat_Sheet.html",
    "Melicher, W., Ur, B., Segreti, S. M., Komanduri, S., Bauer, L., Christin, N., & "
    "Cranor, L. F. (2016). Fast, lean, and accurate: Modeling password guessability using "
    "neural networks. In 25th USENIX Security Symposium.",
    "Gosney, J. (2012). How LinkedIn's password sloppiness hurts us all. Ars Technica.",
    "Hunt, T. (2023). Have I Been Pwned. https://haveibeenpwned.com/",
    "Python Software Foundation. (2024). hashlib — Secure hashes and message digests. "
    "Python 3.12 documentation.",
    "Streamlit Inc. (2024). Streamlit documentation. https://docs.streamlit.io/",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    para_fmt(p, space_before=1, space_after=5, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    r = p.add_run(f"[{i}]  {ref}")
    set_font(r, size=11)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# APPENDICES
# ══════════════════════════════════════════════════════════════
add_heading(doc, "APPENDICES", level=0)

add_heading(doc, "Appendix A — Sample Password File (passwords.txt)", level=1, numbered="A")
add_body(doc, "The following is a representative excerpt from the generated passwords.txt file:")
add_code_block(doc, """
password
123456
qwerty
letmein
dragon
admin
monkey
abc123
iloveyou
12345678
master
sunshine
princess
welcome
login
passw0rd
football
shadow
superman
batman
michael
thomas
soccer
charlie
donald
password1
zxcvbn
trustno1
baseball
hunter
""")
add_caption(doc, "Figure A.1 — Excerpt from data/passwords.txt (30 of ~55 entries shown)")

add_heading(doc, "Appendix B — Sample Hash File (hashes_md5.txt)", level=1, numbered="B")
add_body(doc, "Each line contains an MD5 hash followed by its corresponding plaintext password:")
add_code_block(doc, """
5f4dcc3b5aa765d61d8327deb882cf99: password
e10adc3949ba59abbe56e057f20f883e: 123456
d8578edf8458ce06fbc5bb76a58c5ca4: qwerty
0d107d09f5bbe40cade3de5c71e9e9b7: letmein
8621ffdbc5698829397d97767ac13db3: dragon
21232f297a57a5a743894a0e4a801fc3: admin
d0763edaa9d9bd2a9516280e9044d885: monkey
e99a18c428cb38d5f260853678922e03: abc123
f25a2fc72690b780b2a14e140ef6a9e0: iloveyou
25d55ad283aa400af464c76d713c07ad: 12345678
""")
add_caption(doc, "Figure B.1 — Excerpt from data/hashes_md5.txt (10 of ~55 entries shown)")

add_heading(doc, "Appendix C — Attack Report Summary", level=1, numbered="C")
add_body(doc, "Structured summary of all four attack experiments:")
add_table(doc,
    ["Field", "Dict. MD5", "Dict. SHA-1", "BF MD5", "BF SHA-1"],
    [
        ["Targets",       "30",   "30",   "5",         "5"],
        ["Cracked",       "30",   "30",   "5",         "5"],
        ["Success Rate",  "100%", "100%", "100%",      "100%"],
        ["Attempts",      "30",   "30",   "1,342,359", "1,342,359"],
        ["Time (s)",      "0.0002","0.0002","4.6942",  "4.5450"],
        ["Charset",       "N/A",  "N/A",  "a-z+0-9",  "a-z+0-9"],
        ["Max length",    "N/A",  "N/A",  "5 chars",   "5 chars"],
    ],
    caption="Table C.1 — Full attack results matrix"
)

add_heading(doc, "Appendix D — Application Screenshots", level=1, numbered="D")
add_body(doc, (
    "The following descriptions correspond to the application screenshots visible in the "
    "interactive Streamlit interface of CryptLab:"
))

screenshots = [
    ("D.1", "Overview Page", (
        "The Overview page presents the research context, a KPI dashboard showing 4 hashing "
        "algorithms, 2 attack vectors, 21 ML features, 3 defence layers, and 10 report pages. "
        "The Research Modules section shows three panels: Module I (Vulnerability Assessment), "
        "Module II (Cryptanalysis & AI), and Module III (Defence Architecture)."
    )),
    ("D.2", "Hash Lab — Dataset Generation", (
        "The Hash Lab page allows users to enter a password and observe its MD5 and SHA-1 hashes "
        "in real time, generate a weak password dataset with configurable size, and view the "
        "vulnerability demonstration showing identical hashes for identical passwords."
    )),
    ("D.3", "Attack Simulation — Dictionary Attack Results", (
        "The Attack Simulation page displays a table of cracked passwords alongside their "
        "hash values, with metrics showing 30/30 targets cracked at 100% success rate in "
        "0.0002 seconds. Three comparison bar charts visualise success rate, cracking time, "
        "and hash attempts across attack types."
    )),
    ("D.4", "Cryptanalysis — Entropy Distribution", (
        "The Cryptanalysis page shows a histogram of Shannon entropy values across the password "
        "dataset and a scatter plot of password length versus entropy, revealing the low entropy "
        "characteristic of the weak password set."
    )),
    ("D.5", "Defence System — bcrypt and Argon2 Hashing", (
        "The Defence System page allows users to hash any password with bcrypt or Argon2, "
        "displaying the resulting hash (with embedded salt), computation time, and verification "
        "result. The policy checker displays a checklist of rule compliance for any entered password."
    )),
    ("D.6", "Security Intelligence — Threat Report", (
        "The Security Intelligence page generates a full threat-level panel showing an overall "
        "security score out of 100, estimated crack times under dictionary, brute-force, and "
        "bcrypt-protected scenarios, live hash generation, AI classification, and policy audit."
    )),
]

for fig_id, title, desc in screenshots:
    add_body(doc, f"Figure {fig_id} — {title}")
    p = doc.add_paragraph()
    para_fmt(p, space_before=2, space_after=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    r = p.add_run(desc)
    set_font(r, size=11, italic=True, color=(89, 89, 89))
    p.paragraph_format.left_indent = Cm(0.8)

# ── Save ───────────────────────────────────────────────────────
output_path = "CryptLab_Project_Report.docx"
doc.save(output_path)
print(f"\n✓  Report saved → {output_path}")
print(f"   Chapters: 16 + Appendices A–D")
print(f"   Estimated pages: 35–40")
